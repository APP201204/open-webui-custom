"""Utility for translating .docx files post-generation.

This module provides functionality to translate the text content of
Word documents while preserving formatting, structure, and code blocks.
"""

import logging
from pathlib import Path
from typing import Optional

from open_webui.env import ENABLE_TRANSLATION
from open_webui.utils.translation import translate_batch

log = logging.getLogger(__name__)


async def translate_docx(
    docx_path: str,
    target_lang: str,
    output_path: Optional[str] = None,
) -> str:
    """Translate a .docx file to the target language.

    Walks every text run (paragraphs, headings, bullets, table cells, header/footer),
    collects non-empty text, translates in batch, and writes back.

    Args:
        docx_path: Path to the input .docx file.
        target_lang: Target language code (e.g., 'hi', 'gu', 'es').
        output_path: Optional output path. If None, creates a new file with
            `original__<lang>.docx` suffix.

    Returns:
        Path to the translated .docx file.

    Raises:
        ValueError: If translation is disabled or target language is English.
        FileNotFoundError: If input file doesn't exist.
    """
    if not ENABLE_TRANSLATION:
        raise ValueError("Translation is disabled. Set ENABLE_TRANSLATION=True to enable.")

    if target_lang == 'en':
        log.info("Target language is English, skipping translation")
        return docx_path

    docx_file = Path(docx_path)
    if not docx_file.exists():
        raise FileNotFoundError(f"Input file not found: {docx_path}")

    # Determine output path
    if output_path is None:
        output_path = str(docx_file.parent / f"{docx_file.stem}__{target_lang}{docx_file.suffix}")

    try:
        from docx import Document

        doc = Document(docx_path)

        # Collect all text runs with their references
        text_runs = []

        # Helper to collect text from a collection of runs
        def collect_runs_from_paragraph(paragraph):
            for run in paragraph.runs:
                if run.text and run.text.strip():
                    text_runs.append((run, run.text))

        # Helper to collect text from table cells
        def collect_runs_from_table(table):
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        collect_runs_from_paragraph(paragraph)

        # Collect from paragraphs
        for paragraph in doc.paragraphs:
            collect_runs_from_paragraph(paragraph)

        # Collect from tables
        for table in doc.tables:
            collect_runs_from_table(table)

        # Collect from headers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        collect_runs_from_paragraph(paragraph)
                    for table in header.tables:
                        collect_runs_from_table(table)

            # Collect from footers
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        collect_runs_from_paragraph(paragraph)
                    for table in footer.tables:
                        collect_runs_from_table(table)

        if not text_runs:
            log.info("No text found in document, skipping translation")
            return docx_path

        # Extract text strings for batch translation
        original_texts = [text for _, text in text_runs]

        # Translate in batch
        log.info(f"Translating {len(original_texts)} text runs to {target_lang}")
        translated_texts = await translate_batch(original_texts, target_lang)

        # Write translated text back to runs
        for (run, _), translated in zip(text_runs, translated_texts):
            run.text = translated

        # Save to new path
        doc.save(output_path)
        log.info(f"Translated document saved to: {output_path}")

        return output_path

    except ImportError:
        log.error("python-docx not installed. Install with: pip install python-docx")
        raise ValueError("python-docx is required for docx translation")
    except Exception as e:
        log.error(f"Failed to translate docx file: {e}")
        raise


def get_user_language_from_metadata(metadata: dict) -> Optional[str]:
    """Extract user language from chat metadata.

    Args:
        metadata: Chat metadata dictionary (from form_data['metadata']).

    Returns:
        Language code if set and not English, None otherwise.
    """
    user_language = metadata.get('user_language') if metadata else None
    if user_language and user_language != 'en':
        return user_language
    return None
